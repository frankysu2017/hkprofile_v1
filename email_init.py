#!/usr/bin/env python3
# coding=utf-8
# email_init.py

import email
import re
from bs4 import BeautifulSoup
from dateutil.parser import parse
from docx import Document
import win32com
import win32com.client
import xlrd
from sqlalchemy import Column, String, create_engine, Integer, DateTime
from sqlalchemy.orm import sessionmaker
from sqlalchemy.ext.declarative import declarative_base

wc = win32com.client.constants
Base = declarative_base()


class Email(Base):
    __tablename__ = 'email'

    id = Column(Integer, primary_key=True, autoincrement=True)
    send_name = Column(String(255))
    send_box = Column(String(255))
    receive_name = Column(String(255))
    receive_box = Column(String(255))
    send_time = Column(DateTime)
    subject = Column(String(255))

    def __repr__(self):
        return '<Email(send_name="%s", send_box="%s", receive_name="%s", receive_box="%s", send_time="%s", subject="%s">' \
               % (self.send_name, self.send_box, self.receive_name, self.receive_box, self.send_time, self.subject)





def get_subject(msg):
    s = msg.get('subject')
    if s:
        s_decoded = email.header.decode_header(s)
        return s_decoded[0][0].decode(s_decoded[0][1], 'ignore')
    else:
        return None


def get_sendtime(msg):
    if 'date' in msg:
        return parse(msg.get('date'))
    else:
        return None


def get_sender(msg):
    sname, send_box = email.utils.parseaddr(msg.get('from'))
    if sname:
        sname = email.header.decode_header(sname)
        if isinstance(sname[0][0], bytes):
            send_name = sname[0][0].decode(sname[0][1], 'ignore')
        else:
            send_name = sname[0][0]
    else:
        send_name = None
    return send_name, send_box


def get_receiver(msg):
    rname, receive_box = email.utils.parseaddr(msg.get('to'))
    if rname:
        rname = email.header.decode_header(rname)
        if isinstance(rname[0][0], bytes):
            receive_name = rname[0][0].decode(rname[0][1], 'ignore')
        else:
            receive_name = rname[0][0]
    else:
        receive_name = rname
    return receive_name, receive_box


def get_mainbody(msg):
    mainbody = msg.get_payload(decode=True).strip().decode(msg.get_content_charset(), 'ignore')
    html_flag = re.search('text/(.*)', msg.get_content_type()).group(1).lower()
    if html_flag == 'html':
        prefix, html_text, suffix = re.search('(.*)(<html>.*</html>)(.*)', mainbody, re.DOTALL+re.IGNORECASE).groups()
        html_text = BeautifulSoup(html_text, features="html.parser").body.get_text()
        mainbody = prefix + '\n' + html_text + '\n' + suffix
    mainbody = re.sub('[\n]+', '\n', mainbody)
    return mainbody


def read_excel(filename):
    excel = xlrd.open_workbook(filename)
    attachment_text = ''
    for sheet in excel.sheets():
        for row in sheet.get_rows():
            attachment_text += ','.join([str(cell.value) for cell in row])
        attachment_text += '\n\n\n'
    return attachment_text


attachment_file_path = 'C:\\Users\\junie\\PycharmProjects\\hkprofile_v1\\attachments\\'
def read_word(filename):
    extension_name = filename.split('.')[-1].lower()
    if extension_name == 'doc':
        try:
            wps = win32com.client.gencache.EnsureDispatch('kwps.application')
        except:
            wps = win32com.client.gencache.EnsureDispatch('wps.application')
        else:
            wps = win32com.client.gencache.EnsureDispatch('word.application')
        d = wps.Documents.Open(filename)
        d.SaveAs2(attachment_file_path + 'temp.docx', 12)
        try:
            wps.Documents.Close()
            wps.Documents.Close(wc.wdDoNotSaveChanges)
            wps.Quit
        except:
            pass
        try:
            doc = Document(attachment_file_path + r'temp.docx')
        except:
            pass
    else:
        try:
            doc = Document(filename)
        except:
            pass
    attachemnt_text = ''
    for each in doc.paragraphs:
        attachemnt_text += each.text + '\n'
    return attachemnt_text


def get_attachment(msg):
    attachemnt_name = msg.get_filename()
    if attachemnt_name:
        attachemnt_name = email.header.decode_header(attachemnt_name)
        attachemnt_name = attachemnt_name[0][0].decode(attachemnt_name[0][1], 'ignore')
        attachemnt_content = msg.get_payload(decode=True).strip()
        with open(attachment_file_path + attachemnt_name, 'wb') as f:
            f.write(attachemnt_content)
        extension_name = attachemnt_name.split('.')[-1].lower()
        if extension_name in ['doc', 'docx']:
            attachment_text = read_word(attachment_file_path + attachemnt_name)
            return attachemnt_name, attachment_text
        elif extension_name in ['xls', 'xlsx']:
            attachment_text = read_excel(attachment_file_path + attachemnt_name)
            return attachemnt_name, attachment_text
        else:
            return None, None


def email_init(emlfile):
    print(emlfile)
    with open(emlfile, 'r') as eml:
        msg = email.message_from_file(eml)
        subject = get_subject(msg)
        send_time = get_sendtime(msg)
        send_name, send_box = get_sender(msg)
        receive_name, receive_box = get_receiver(msg)
        attachment_file, attachment_text = None, None
        if msg.is_multipart():
            print('it is multipart:\n')
            mail_content = ''
            for part in msg.get_payload():
                if part.get_content_maintype() == 'text':
                    part_content = get_mainbody(part)
                else:
                    part_content = ''
                    attachment_file, attachment_text = get_attachment(part)
                mail_content += part_content
        else:
            print('it is text')
            mail_content = get_mainbody(msg)
        print('{}, {}, {}, {}, {}, {}'.format(send_name, send_box, receive_name, receive_box, send_time, subject))
        print('mail_content: {}'.format(mail_content))
        if attachment_file:
            print('attachment file name: {}'.format(attachment_file))
            print('attachment text: {}'.format(attachment_text))
    engine = create_engine('sqlite:///foo.db')
    Session = sessionmaker(bind=engine)
    session = Session()
    e = Email(send_name=send_name, send_box=send_box, receive_name=receive_name, receive_box=receive_box, send_time=send_time, subject=subject)
    session.add(e)
    session.commit()


if __name__ == '__main__':
    #print('send_name, send_box, receive_name, receive_box, send_time, subject')
    email_init(r'./email/01.eml')
    email_init(r'./email/02.eml')
    email_init(r'./email/03.eml')
    email_init(r'./email/04.eml')
    #read_excel(r'./test.xlsx')